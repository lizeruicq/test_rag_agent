"""
文档读取器 - 支持多种文档格式
"""
import os
from typing import List, Dict, Any
from abc import ABC, abstractmethod


class DocumentReader(ABC):
    """文档读取器基类"""

    @abstractmethod
    def read(self, file_path: str) -> List[Dict[str, Any]]:
        """
        读取文档并返回分块内容

        Args:
            file_path: 文件路径

        Returns:
            包含 content 和 source 的字典列表
        """
        raise NotImplementedError("Subclasses must implement the read() method")

    def _chunk_text(self, text: str, chunk_size: int = 2000, overlap: int = 200) -> List[str]:
        """
        将文本分块处理

        Args:
            text: 输入文本
            chunk_size: 每块大小
            overlap: 块之间的重叠大小

        Returns:
            文本块列表
        """
        if not text:
            return [""]

        if chunk_size <= 0:
            raise ValueError("chunk_size must be positive")

        # Ensure overlap is smaller than chunk_size to make progress
        overlap = max(0, min(overlap, chunk_size - 1))
        step = chunk_size - overlap

        chunks: List[str] = []
        start = 0
        text_len = len(text)
        while start < text_len:
            end = min(start + chunk_size, text_len)
            chunks.append(text[start:end])
            start += step

        return chunks


class TxtReader(DocumentReader):
    """纯文本文件读取器"""

    def read(self, file_path: str) -> List[Dict[str, Any]]:
        """读取 .txt 文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            chunks = self._chunk_text(content)
            return [{"content": chunk, "source": file_path} for chunk in chunks]
        except Exception as e:
            raise ValueError(f"Failed to read TXT file {file_path}: {str(e)}")


class PdfReader(DocumentReader):
    """PDF 文件读取器"""

    def read(self, file_path: str) -> List[Dict[str, Any]]:
        """读取 .pdf 文件"""
        try:
            import PyPDF2

            documents = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        chunks = self._chunk_text(text)
                        for chunk in chunks:
                            documents.append({
                                "content": chunk,
                                "source": f"{file_path} (Page {page_num + 1})"
                            })

            return documents if documents else [{"content": "", "source": file_path}]
        except Exception as e:
            raise ValueError(f"Failed to read PDF file {file_path}: {str(e)}")


class DocxReader(DocumentReader):
    """Word 文档读取器"""

    def read(self, file_path: str) -> List[Dict[str, Any]]:
        """读取 .docx 文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_content = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # 也处理表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_content.append(row_text)

            full_text = "\n".join(text_content)
            chunks = self._chunk_text(full_text)
            return [{"content": chunk, "source": file_path} for chunk in chunks]
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file {file_path}: {str(e)}")


class ExcelReader(DocumentReader):
    """Excel 文件读取器"""

    def read(self, file_path: str) -> List[Dict[str, Any]]:
        """读取 .xlsx 和 .xls 文件"""
        try:
            import pandas as pd

            documents = []

            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # 转换为文本格式
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string()

                chunks = self._chunk_text(sheet_text)
                for chunk in chunks:
                    documents.append({
                        "content": chunk,
                        "source": f"{file_path} (Sheet: {sheet_name})"
                    })

            return documents if documents else [{"content": "", "source": file_path}]
        except Exception as e:
            raise ValueError(f"Failed to read Excel file {file_path}: {str(e)}")


def get_reader_for_file(file_path: str) -> DocumentReader:
    """
    根据文件扩展名返回相应的读取器

    Args:
        file_path: 文件路径

    Returns:
        对应的 DocumentReader 实例

    Raises:
        ValueError: 如果文件类型不支持
    """
    ext = os.path.splitext(file_path)[1].lower()

    readers = {
        '.txt': TxtReader,
        '.pdf': PdfReader,
        '.docx': DocxReader,
        '.xlsx': ExcelReader,
        '.xls': ExcelReader,
    }

    if ext not in readers:
        supported = ', '.join(readers.keys())
        raise ValueError(f"不支持的文件类型: {ext}。支持的类型: {supported}")

    return readers[ext]()